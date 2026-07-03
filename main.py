"""
主入口：批量处理 Excel → 生成 SEO 长文 → 生图 → 发 WordPress 文章

每次运行处理 DAILY_LIMIT 条；跑过的会在 input.xlsx 里打标 + 填色：
  - 黄色  done @ time          = 成功，下次跳过
  - 红色  error xN @ time      = 失败 N 次（默认下次也跳过；--retry-errors 时才会重试）
  - 灰色  skipped: ...         = 数据校验失败（缺字段等），跳过

正文字数低于 MIN_WORDS 会自动重试，最多 MAX_WORD_RETRIES 次。

⚠️  跑之前请关闭 Excel 中的 input.xlsx，否则脚本无法写回。

注：屏幕输出全部为英文，避免 Windows cmd 中文乱码；
   屏幕和日志由 Python 内置 tee 同时写，不依赖 PowerShell Tee-Object。

用法：
    python main.py
    python main.py --no-wp
    python main.py --limit 1
    python main.py --retry-errors        # 重试历史失败的行
"""
import argparse
import os
import re
import sys
import time
from datetime import datetime
from typing import List, Optional, Tuple

import docx
import markdown as md_lib
import pandas as pd
from bs4 import BeautifulSoup
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

from config import (
    DEFAULT_LANGUAGE,
    IMAGE_ENABLED,
    INPUT_EXCEL,
    OPENROUTER_API_KEY,
    OUTPUT_DIR,
    PAUSE_EVERY_N_ROWS,
    TAVILY_API_KEY,
    WORD_COUNT_MAX,
    WORD_COUNT_MIN,
    WP_APP_PASSWORD,
    WP_URL,
    WP_USERNAME,
)
from image_generator import generate_image
from workflow import ListicleWorkflow
from wordpress_client import WordPressClient


# ============================================================
# 全局常量
# ============================================================
DAILY_LIMIT = 10            # 每次运行处理多少条
MAX_FAIL = 3                # --retry-errors 时,失败几次后放弃
MIN_WORDS = 1000            # 正文最低字数（不达标会重试）
MAX_WORD_RETRIES = 2        # 字数不足时的重试次数（不含首次）
WP_POST_STATUS = "publish"  # WordPress 发布状态: "publish"=直接发布, "draft"=草稿

YELLOW_FILL = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
RED_FILL    = PatternFill(start_color="F8CBAD", end_color="F8CBAD", fill_type="solid")
GRAY_FILL   = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")


# ============================================================
# Tee Logger —— 屏幕 + 日志同时写
# ============================================================
class _Tee:
    def __init__(self, *streams):
        self.streams = streams

    def write(self, data):
        for s in self.streams:
            try:
                s.write(data)
            except Exception:
                pass
        self.flush()

    def flush(self):
        for s in self.streams:
            try:
                s.flush()
            except Exception:
                pass

    def isatty(self):
        try:
            return self.streams[0].isatty()
        except Exception:
            return False

    def fileno(self):
        return self.streams[0].fileno()


def _setup_tee_logging():
    """同时把 stdout/stderr 写到屏幕和当天日志文件"""
    log_dir = "logs"
    os.makedirs(log_dir, exist_ok=True)
    log_path = os.path.join(log_dir, f"run_{datetime.now().strftime('%Y%m%d')}.log")
    log_file = open(log_path, "a", encoding="utf-8", buffering=1)
    log_file.write(f"\n===== Started {datetime.now():%Y-%m-%d %H:%M:%S} =====\n")
    sys.stdout = _Tee(sys.__stdout__, log_file)
    sys.stderr = _Tee(sys.__stderr__, log_file)
    return log_file, log_path


# ============================================================
# 工具函数
# ============================================================
def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", name)[:100].strip()


def slugify_keyword(text: str) -> str:
    """
    把主关键词转成 WP URL slug。
    "Metatrader 5 Broker"        -> "metatrader-5-broker"
    "best forex trading app"     -> "best-forex-trading-app"
    "EUR/USD analysis & tips!"   -> "eur-usd-analysis-tips"
    "外汇 入门 指南"              -> "外汇-入门-指南"  (中文保留)
    长度上限 80 字符,避免 WP 内部再被截断。
    """
    if not text:
        return ""
    s = str(text).strip().lower()
    # 把所有"非字母/数字/中日韩字符"都替成 -
    s = re.sub(r"[^\w\u4e00-\u9fff\u3040-\u309f\u30a0-\u30ff]+", "-", s, flags=re.UNICODE)
    # 折叠多个 -
    s = re.sub(r"-+", "-", s)
    s = s.strip("-")
    return s[:80]


def parse_csv_field(val) -> List[str]:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return []
    s = str(val).strip()
    if not s or s.lower() == "nan":
        return []
    return [x.strip() for x in s.replace("，", ",").split(",") if x.strip()]


def count_words(text: str, language: str = "English") -> int:
    if not text:
        return 0
    lang = language.lower()
    if any(k in lang for k in ["chinese", "中文", "zh"]):
        return len(re.sub(r"\s", "", text))
    return len(text.split())


# ============================================================
# input.xlsx 进度管理
# ============================================================
def _parse_status(val, retry_errors: bool = False) -> Tuple[bool, int]:
    """返回 (是否应跳过, 已失败次数)

    默认行为(retry_errors=False):
        - done / skipped → 跳过
        - error xN       → 跳过(无论 N 是多少)
        - 空 / 其他      → 不跳过

    retry_errors=True 时:
        - done / skipped → 跳过
        - error xN       → N >= MAX_FAIL 才跳过(允许重试)
        - 空 / 其他      → 不跳过
    """
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return False, 0
    s = str(val).strip().lower()
    if not s:
        return False, 0
    if s.startswith("done") or s.startswith("skipped"):
        return True, 0
    m = re.match(r"error\s+x(\d+)", s)
    if m:
        cnt = int(m.group(1))
        if not retry_errors:
            return True, cnt
        return cnt >= MAX_FAIL, cnt
    return False, 0


def ensure_status_column(excel_path: str) -> int:
    wb = load_workbook(excel_path)
    ws = wb.active
    headers = [c.value for c in ws[1]]
    if "status" in headers:
        col = headers.index("status") + 1
        wb.close()
        return col
    col = len(headers) + 1
    ws.cell(row=1, column=col, value="status")
    wb.save(excel_path)
    wb.close()
    return col


def mark_input_row(excel_path: str, excel_row: int, status_col: int,
                   status_text: str, fill: PatternFill):
    try:
        wb = load_workbook(excel_path)
        ws = wb.active
        ws.cell(row=excel_row, column=status_col, value=status_text)
        for c in range(1, ws.max_column + 1):
            ws.cell(row=excel_row, column=c).fill = fill
        wb.save(excel_path)
        wb.close()
    except Exception as e:
        print(f"     [WARN] Failed to write back to input.xlsx: {e} (is it open in Excel?)")


# ============================================================
# Markdown -> Word
# ============================================================
def create_word_document(markdown_text: str, filepath: str) -> Optional[str]:
    try:
        doc = docx.Document()

        def add_formatted_runs(paragraph, text):
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    paragraph.add_run(part[2:-2]).bold = True
                elif part:
                    paragraph.add_run(part)

        lines = markdown_text.strip().split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if line.startswith("|") and line.endswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1

                if table_lines:
                    data_rows = [l for l in table_lines if not re.match(r"^[|\s:-]+$", l)]
                    if not data_rows:
                        continue
                    header_cols = [c.strip() for c in data_rows[0].split("|")[1:-1]]
                    table = doc.add_table(rows=1, cols=len(header_cols))
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    for j, h in enumerate(header_cols):
                        p = hdr_cells[j].paragraphs[0]
                        p.clear()
                        add_formatted_runs(p, h)
                        hdr_cells[j].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    for row_text in data_rows[1:]:
                        cells_text = [c.strip() for c in row_text.split("|")[1:-1]]
                        row_cells = table.add_row().cells
                        for j, ct in enumerate(cells_text):
                            if j < len(row_cells):
                                p = row_cells[j].paragraphs[0]
                                p.clear()
                                add_formatted_runs(p, ct)
                continue

            if not line:
                i += 1
                continue

            if line.startswith("# "):
                add_formatted_runs(doc.add_heading(level=1), line[2:].strip())
            elif line.startswith("## "):
                add_formatted_runs(doc.add_heading(level=2), line[3:].strip())
            elif line.startswith("### "):
                add_formatted_runs(doc.add_heading(level=3), line[4:].strip())
            elif line.startswith("#### "):
                add_formatted_runs(doc.add_heading(level=4), line[5:].strip())
            elif line.startswith("* ") or line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_runs(p, line[2:].strip())
            else:
                p = doc.add_paragraph()
                add_formatted_runs(p, line)
            i += 1

        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        doc.save(filepath)
        return filepath
    except Exception as e:
        print(f"     [ERR] Word generation failed: {e}")
        return None


# ============================================================
# Markdown -> HTML
# ============================================================
def markdown_to_html_for_wp(
    markdown_text: str,
    inline_image_url: Optional[str] = None,
    inline_alt: str = "",
) -> Tuple[str, str]:
    title_match = re.search(r"^#\s+(.+)$", markdown_text, re.MULTILINE)
    title_text = title_match.group(1).strip() if title_match else ""
    # 去掉模型可能在标题结尾带上的字数标注，如 "[~12 từ]" / "[~150 words]" / "[12 字]" 等。
    # 规则：结尾的方括号里若含数字，就整段方括号连同前面的空格一起删掉。
    title_text = re.sub(r"\s*[\[\(（【][^\]\)）】]*\d+[^\]\)）】]*[\]\)）】]\s*$", "", title_text).strip()
    body_md = re.sub(r"^#\s+.+\n?", "", markdown_text, count=1, flags=re.MULTILINE)

    html = md_lib.markdown(body_md, extensions=["tables", "fenced_code"])

    if inline_image_url:
        soup = BeautifulSoup(html, "html.parser")
        first_h2 = soup.find("h2")
        if first_h2:
            fig = soup.new_tag("figure", attrs={"class": "wp-block-image size-large aligncenter"})
            img = soup.new_tag(
                "img",
                src=inline_image_url,
                alt=inline_alt or "article illustration",
                loading="lazy",
            )
            fig.append(img)
            if inline_alt:
                cap = soup.new_tag("figcaption")
                cap.string = inline_alt
                fig.append(cap)
            first_h2.insert_after(fig)
            html = str(soup)

    return title_text, html


# ============================================================
# 单条处理流程
# ============================================================
def process_row(
    row_data: dict,
    row_index: int,
    language: str,
    output_dir: str,
    workflow: ListicleWorkflow,
    wp_client: Optional[WordPressClient],
) -> dict:
    start = time.time()
    idx = row_index + 1
    print(f"\n{'='*60}")
    print(f">> [{idx}] {row_data['main_keyword']} | {language} | {row_data['wordcounts']} words")
    print(f"{'='*60}")

    result = {
        "row_index": idx,
        "language": language,
        **row_data,
        "status": "Failed",
        "process_time": "0s",
        "outline": "",
        "optimized_article": "",
        "seo_title": "",
        "seo_description": "",
        "word_count": 0,
        "word_filename": "",
        "image_file": "",
        "wp_post_id": "",
        "wp_post_link": "",
    }
    result["categories"] = ", ".join(row_data.get("categories", []))
    result["tags"] = ", ".join(row_data.get("tags", []))
    result["cost_usd"] = 0.0
    workflow.row_cost = 0.0  # 重置本篇 OpenRouter 花费计数

    try:
        # -------- 1. 搜索 --------
        print("[1] Tavily search (main keyword)...")
        main_res = workflow.tavily_search(row_data["main_keyword"], max_results=5)
        time.sleep(1.5)

        print("[2] Tavily search (secondary keyword)...")
        sec_res = workflow.tavily_search(row_data["secondary_keyword"], max_results=5)
        time.sleep(1.5)

        # -------- 2. 大纲 --------
        print("[3] Generating outline...")
        outline = workflow.generate_outline(
            row_data["main_keyword"],
            row_data["secondary_keyword"],
            row_data["topic"],
            row_data["wordcounts"],
            row_data["specific"],
            main_res,
            sec_res,
            language,
        )
        if not outline or outline.startswith("LLM call error"):
            raise RuntimeError(f"Outline generation failed: {outline}")
        result["outline"] = outline
        time.sleep(1.5)

        # -------- 3. 正文（带字数校验重试） --------
        print("[4] Writing article...")
        article = ""
        word_count = 0
        last_err = None

        for attempt in range(1, MAX_WORD_RETRIES + 2):
            specific_for_call = row_data["specific"]
            if attempt > 1:
                specific_for_call = (
                    (specific_for_call + " " if specific_for_call else "")
                    + f"IMPORTANT: The article MUST be at least {MIN_WORDS} words. "
                    f"Previous draft only had {word_count} words; please expand each section "
                    f"with more detail, examples, and analysis."
                )

            article = workflow.write_article(
                row_data["main_keyword"],
                row_data["secondary_keyword"],
                row_data["topic"],
                row_data["wordcounts"],
                specific_for_call,
                outline,
                main_res,
                sec_res,
                language,
            )

            if not article or article.startswith("LLM call error"):
                last_err = article
                print(f"     [WARN] Attempt {attempt} failed: {article}")
                if attempt <= MAX_WORD_RETRIES:
                    time.sleep(3)
                    continue
                raise RuntimeError(f"Article generation failed: {last_err}")

            word_count = count_words(article, language)
            if word_count >= MIN_WORDS:
                print(f"     [OK] Article {word_count} words (attempt {attempt})")
                break

            print(f"     [WARN] Attempt {attempt} only {word_count} words (min {MIN_WORDS})")
            if attempt <= MAX_WORD_RETRIES:
                print(f"     [INFO] Retrying...")
                time.sleep(3)
            else:
                print(f"     [WARN] Retried {MAX_WORD_RETRIES} times, still short. Using current version.")

        result["optimized_article"] = article
        result["word_count"] = word_count
        time.sleep(1.5)

        # -------- 4. SEO --------
        print("[5] Generating SEO metadata...")
        seo = workflow.generate_seo_meta(article, row_data["main_keyword"], language)
        result["seo_title"] = seo["seo_title"]
        result["seo_description"] = seo["seo_description"]

        # -------- 5. 图片 --------
        image_path = None
        if IMAGE_ENABLED:
            print("[6] Generating image prompt...")
            img_prompt = workflow.generate_image_prompt(article, row_data["main_keyword"])

            img_dir = os.path.join(output_dir, "images")
            os.makedirs(img_dir, exist_ok=True)
            slug = sanitize_filename(row_data["main_keyword"]).replace(" ", "_").lower() or "post"

            print("[7] Generating image (Nano Banana)...")
            image_path = generate_image(
                img_prompt,
                os.path.join(img_dir, f"{idx}_{slug}.png"),
            )
            if image_path:
                result["image_file"] = image_path
                print(f"     [OK] Image: {image_path}")

        # -------- 6. Word --------
        print("[8] Generating Word doc...")
        word_dir = os.path.join(output_dir, "articles")
        word_name = f"{idx}_{sanitize_filename(row_data['main_keyword'])}.docx"
        word_path = os.path.join(word_dir, word_name)
        if create_word_document(article, word_path):
            result["word_filename"] = word_path
            print(f"     [OK] Word: {word_path}")

        # -------- 7. WordPress --------
        if wp_client:
            print(f"[9] Publishing to WordPress (status={WP_POST_STATUS})...")
            media = None
            if image_path:
                media = wp_client.upload_media(
                    image_path,
                    alt_text=row_data["main_keyword"],
                )
                if media:
                    print(f"     [OK] Image uploaded: ID={media['id']}")

            title_text, html_body = markdown_to_html_for_wp(
                article,
                inline_image_url=media["url"] if media else None,
                inline_alt=row_data["main_keyword"],
            )
            post_title = title_text or seo["seo_title"]
            post_slug = slugify_keyword(row_data["main_keyword"])
            if post_slug:
                print(f"     URL slug: {post_slug}")

            cat_ids = wp_client.resolve_categories(row_data.get("categories", []))
            tag_ids = wp_client.resolve_tags(row_data.get("tags", []))
            if cat_ids:
                print(f"     Category IDs: {cat_ids}")
            if tag_ids:
                print(f"     Tag IDs: {tag_ids}")

            post = wp_client.create_post(
                title=post_title,
                content_html=html_body,
                excerpt=seo["seo_description"],
                featured_media=media["id"] if media else None,
                categories=cat_ids,
                tags=tag_ids,
                status=WP_POST_STATUS,
                slug=post_slug,
            )
            if post and post.get("id"):
                result["wp_post_id"] = post["id"]
                result["wp_post_link"] = post["link"]
                print(f"     [OK] Post {WP_POST_STATUS}: {post['link']}")
            else:
                print(f"     [WARN] Post creation failed (status={WP_POST_STATUS})")

        result["status"] = "Success"
        print(f"[OK] [{idx}] Done")

    except Exception as e:
        import traceback
        traceback.print_exc()
        result["status"] = f"Failed: {e}"
        print(f"[ERR] [{idx}] Failed: {e}")

    result["process_time"] = f"{time.time() - start:.1f}s"
    result["cost_usd"] = round(workflow.row_cost, 6)
    print(f"[$] [{idx}] OpenRouter 本篇花费: ${workflow.row_cost:.4f} "
          f"| 累计: ${workflow.total_cost:.4f}")
    return result


# ============================================================
# 主流程
# ============================================================
def main():
    parser = argparse.ArgumentParser(
        description="SEO Blog Generator - with Nano Banana images & WordPress publishing"
    )
    parser.add_argument("--input", "-i", default=INPUT_EXCEL, help="Input Excel path")
    parser.add_argument("--language", "-l", default=DEFAULT_LANGUAGE, help="Article language")
    parser.add_argument("--output", "-o", default=OUTPUT_DIR, help="Output directory")
    parser.add_argument("--no-wp", action="store_true", help="Skip WordPress publishing")
    parser.add_argument("--limit", type=int, default=DAILY_LIMIT,
                        help=f"Max records this run (default {DAILY_LIMIT})")
    parser.add_argument("--retry-errors", action="store_true",
                        help="Retry rows with previous errors (default: skip them)")
    args = parser.parse_args()

    # ------- 开启 tee 日志 -------
    log_file, log_path = _setup_tee_logging()
    print(f"[INFO] Log: {log_path}")
    if args.retry_errors:
        print("[INFO] --retry-errors enabled: previous error rows will be retried")

    try:
        # ------- 校验 -------
        if not TAVILY_API_KEY or not OPENROUTER_API_KEY:
            print("[ERR] Please set TAVILY_API_KEY and OPENROUTER_API_KEY in config.py")
            sys.exit(1)

        if not os.path.exists(args.input):
            print(f"[ERR] Input file not found: {args.input}")
            sys.exit(1)

        os.makedirs(args.output, exist_ok=True)

        # ------- 读 Excel + 加 status 列 -------
        try:
            status_col = ensure_status_column(args.input)
            df = pd.read_excel(args.input)
            if "status" not in df.columns:
                df["status"] = ""
            print(f"[OK] Loaded {len(df)} records: {args.input}")
        except Exception as e:
            print(f"[ERR] Failed to read Excel: {e}")
            sys.exit(1)

        # ------- 筛选本次要跑的 -------
        pending_indices = []
        done_count = 0
        for i, row in df.iterrows():
            skip, _ = _parse_status(row.get("status"), retry_errors=args.retry_errors)
            if skip:
                done_count += 1
                continue
            pending_indices.append(i)
            if len(pending_indices) >= args.limit:
                break

        remaining_total = len(df) - done_count
        print(f"[INFO] This run: {len(pending_indices)} | Done/skipped: {done_count} | Remaining: {remaining_total}")

        if not pending_indices:
            print("[OK] All records processed. Nothing to do.")
            return

        # ------- WP 客户端 -------
        wp_client = None
        if args.no_wp:
            print("[INFO] --no-wp enabled, local-only mode")
        else:
            if not all([WP_URL, WP_USERNAME, WP_APP_PASSWORD]) or WP_URL == "https://your-site.com":
                print("[WARN] WordPress config incomplete (config.py), skipping publish")
            else:
                wp_client = WordPressClient()
                print(f"[OK] WordPress ready: {WP_URL} (post status: {WP_POST_STATUS})")

        # ------- 批量处理 -------
        workflow = ListicleWorkflow()
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        results: List[dict] = []
        results_excel = os.path.join(args.output, f"results_{ts}.xlsx")

        for n, i in enumerate(pending_indices):
            row = df.iloc[i]
            excel_row = i + 2
            _, prev_fail = _parse_status(row.get("status"), retry_errors=args.retry_errors)

            # 预校验
            try:
                wc_raw = row.get("wordcounts")
                if pd.isna(wc_raw):
                    raise ValueError("wordcounts is empty")
                wc = int(wc_raw)
                if wc < WORD_COUNT_MIN:
                    print(f"     [INFO] Excel row {excel_row} word count {wc} below min, raised to {WORD_COUNT_MIN}")
                    wc = WORD_COUNT_MIN
                elif wc > WORD_COUNT_MAX:
                    print(f"     [INFO] Excel row {excel_row} word count {wc} above max, capped to {WORD_COUNT_MAX}")
                    wc = WORD_COUNT_MAX

                row_data = {
                    "main_keyword": str(row["main_keyword"]).strip(),
                    "secondary_keyword": str(row["secondary_keyword"]).strip(),
                    "topic": str(row["topic"]).strip(),
                    "wordcounts": wc,
                    "specific": str(row.get("specific", "")).strip() if pd.notna(row.get("specific", "")) else "",
                    "categories": ["Giáo dục"],  # 固定分类为 Giáo dục（越南语 education，忽略 Excel 的 category 列）
                    "tags": parse_csv_field(row.get("tags")),
                }
                if not all(row_data[k] for k in ["main_keyword", "secondary_keyword", "topic"]):
                    raise ValueError("main_keyword / secondary_keyword / topic cannot be empty")

            except Exception as e:
                print(f"[ERR] Skipping Excel row {excel_row}: {e}")
                mark_input_row(args.input, excel_row, status_col,
                               f"skipped: {e}", GRAY_FILL)
                continue

            # 跑
            r = process_row(row_data, i, args.language, args.output, workflow, wp_client)
            results.append(r)

            # 回写 input.xlsx
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
            if r["status"] == "Success":
                mark_input_row(args.input, excel_row, status_col,
                               f"done @ {now_str}", YELLOW_FILL)
            else:
                new_fail = prev_fail + 1
                err_msg = str(r["status"]).replace("Failed:", "").strip()[:60]
                mark_input_row(args.input, excel_row, status_col,
                               f"error x{new_fail} @ {now_str} | {err_msg}", RED_FILL)
                if new_fail >= MAX_FAIL:
                    print(f"     [INFO] Row failed {new_fail} times.")

            # 即时保存结果
            try:
                pd.DataFrame(results).to_excel(results_excel, index=False)
            except Exception as e:
                print(f"     [WARN] Failed to save results Excel: {e}")

            if (n + 1) % PAUSE_EVERY_N_ROWS == 0 and n + 1 < len(pending_indices):
                print(f"\n[PAUSE] Processed {n+1} records, pausing 15s...\n")
                time.sleep(15)

        # ------- 总结 -------
        ok = sum(1 for r in results if r["status"] == "Success")
        wp_ok = sum(1 for r in results if r.get("wp_post_id"))
        print(f"\n{'='*60}")
        print(f"[DONE] Run finished")
        print(f"   Total processed: {len(results)}")
        print(f"   Success:         {ok}")
        print(f"   Posts {WP_POST_STATUS}:   {wp_ok}")
        print(f"   OpenRouter cost: ${workflow.total_cost:.4f}"
              + (f" (≈ ${workflow.total_cost / len(results):.4f}/篇)" if results else ""))
        print(f"   Results Excel:   {results_excel}")
        print(f"   Word dir:        {os.path.join(args.output, 'articles')}")
        print(f"   Image dir:       {os.path.join(args.output, 'images')}")
        print(f"{'='*60}")

    finally:
        try:
            log_file.write(f"===== Finished {datetime.now():%Y-%m-%d %H:%M:%S} =====\n\n")
            log_file.flush()
            log_file.close()
        except Exception:
            pass


if __name__ == "__main__":
    main()