"""
纯文本工具函数（不依赖 pandas）——供 Web 后台复用，避免启动就加载 pandas。

内容与 main.py 里同名函数一致；main.py(命令行版)仍用自己的那份，互不影响。
"""
import os
import re
from typing import List, Optional, Tuple

import docx
import markdown as md_lib
from bs4 import BeautifulSoup


def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", name)[:100].strip()


def slugify_keyword(text: str) -> str:
    """把主关键词转成 WP URL slug（中文保留），长度上限 80。"""
    if not text:
        return ""
    s = str(text).strip().lower()
    s = re.sub(r"[^\w一-鿿぀-ゟ゠-ヿ]+", "-", s, flags=re.UNICODE)
    s = re.sub(r"-+", "-", s)
    s = s.strip("-")
    return s[:80]


def parse_csv_field(val) -> List[str]:
    # 不依赖 pandas 的 NaN 判断：float 的 NaN 满足 val != val
    if val is None or (isinstance(val, float) and val != val):
        return []
    s = str(val).strip()
    if not s or s.lower() == "nan":
        return []
    return [x.strip() for x in s.replace("，", ",").split(",") if x.strip()]


def count_words(text: str, language: str = "English") -> int:
    if not text:
        return 0
    lang = language.lower()
    if any(k in lang for k in ["chinese", "中文", "zh"]):
        return len(re.sub(r"\s", "", text))
    return len(text.split())


def create_word_document(markdown_text: str, filepath: str) -> Optional[str]:
    try:
        doc = docx.Document()

        def add_formatted_runs(paragraph, text):
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    paragraph.add_run(part[2:-2]).bold = True
                elif part:
                    paragraph.add_run(part)

        lines = markdown_text.strip().split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if line.startswith("|") and line.endswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1

                if table_lines:
                    data_rows = [l for l in table_lines if not re.match(r"^[|\s:-]+$", l)]
                    if not data_rows:
                        continue
                    header_cols = [c.strip() for c in data_rows[0].split("|")[1:-1]]
                    table = doc.add_table(rows=1, cols=len(header_cols))
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    for j, h in enumerate(header_cols):
                        p = hdr_cells[j].paragraphs[0]
                        p.clear()
                        add_formatted_runs(p, h)
                        hdr_cells[j].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    for row_text in data_rows[1:]:
                        cells_text = [c.strip() for c in row_text.split("|")[1:-1]]
                        row_cells = table.add_row().cells
                        for j, ct in enumerate(cells_text):
                            if j < len(row_cells):
                                p = row_cells[j].paragraphs[0]
                                p.clear()
                                add_formatted_runs(p, ct)
                continue

            if not line:
                i += 1
                continue

            if line.startswith("# "):
                add_formatted_runs(doc.add_heading(level=1), line[2:].strip())
            elif line.startswith("## "):
                add_formatted_runs(doc.add_heading(level=2), line[3:].strip())
            elif line.startswith("### "):
                add_formatted_runs(doc.add_heading(level=3), line[4:].strip())
            elif line.startswith("#### "):
                add_formatted_runs(doc.add_heading(level=4), line[5:].strip())
            elif line.startswith("* ") or line.startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_runs(p, line[2:].strip())
            else:
                p = doc.add_paragraph()
                add_formatted_runs(p, line)
            i += 1

        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        doc.save(filepath)
        return filepath
    except Exception as e:
        print(f"     [ERR] Word generation failed: {e}")
        return None


def markdown_to_html_for_wp(
    markdown_text: str,
    inline_image_url: Optional[str] = None,
    inline_alt: str = "",
) -> Tuple[str, str]:
    title_match = re.search(r"^#\s+(.+)$", markdown_text, re.MULTILINE)
    title_text = title_match.group(1).strip() if title_match else ""
    title_text = re.sub(r"\s*[\[\(（【][^\]\)）】]*\d+[^\]\)）】]*[\]\)）】]\s*$", "", title_text).strip()
    body_md = re.sub(r"^#\s+.+\n?", "", markdown_text, count=1, flags=re.MULTILINE)

    html = md_lib.markdown(body_md, extensions=["tables", "fenced_code"])

    if inline_image_url:
        soup = BeautifulSoup(html, "html.parser")
        first_h2 = soup.find("h2")
        if first_h2:
            fig = soup.new_tag("figure", attrs={"class": "wp-block-image size-large aligncenter"})
            img = soup.new_tag(
                "img",
                src=inline_image_url,
                alt=inline_alt or "article illustration",
                loading="lazy",
            )
            fig.append(img)
            if inline_alt:
                cap = soup.new_tag("figcaption")
                cap.string = inline_alt
                fig.append(cap)
            first_h2.insert_after(fig)
            html = str(soup)

    return title_text, html
